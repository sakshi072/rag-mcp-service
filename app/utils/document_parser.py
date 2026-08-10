"""
Advanced Document Parser

Handles:
- PDF files (pypdf + pdfplumber fallback)
- DOCX files (python-docx)
- Plain text files
- Markdown files
"""

import re
from io import BytesIO
from typing import Dict, Any
import pypdf
import pdfplumber
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Parse different document formats and extract text + metadata.

    Supported formats: PDF, DOCX, TXT, MD
    """

    @staticmethod
    def parse_pdf(file_data: bytes) -> Dict[str, Any]:
        """
        Parse PDF file to extract text and metadata.

        Uses pypdf for speed, falls back to pdfplumber for complex PDFs.

        Args:
            file_data: PDF file as bytes

        Returns:
            Dictionary with 'text' and 'metadata'
        """
        try:
            # Try pypdf first (faster)
            pdf_file = BytesIO(file_data)
            reader = pypdf.PdfReader(pdf_file)

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    # Add Page marker for better context
                    text_parts.append(f"[Page {page_num}]\n{text}")

            extracted_text = "\n\n".join(text_parts)

            # If extraction is poor, try pdfplumber
            if len(extracted_text.strip()) < 100:
                logger.warning("   pypdf extraction insufficient, trying pdfplumber...")
                pdf_file = BytesIO(file_data)
                with pdfplumber.open(pdf_file) as pdf:
                    text_parts = []
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"[Page {page_num}]\n{text}")
                    extracted_text = "\n\n".join(text_parts)

            # Extract metadata
            metadata = {
                "pages": len(reader.pages),
                "format": "pdf",
                "file_size": len(file_data)
            }

            # Try to get PDF metadata
            if reader.metadata:
                metadata["author"] = reader.metadata.get("/Author", "Unknown")
                metadata["title"] = reader.metadata.get("/Title", "Untitled")
                metadata["subject"] = reader.metadata.get("/Subject", "")

            cleaned_data = DocumentParser.clean_text(extracted_text)

            return {
                "text": cleaned_data,
                "metadata": metadata
            }
        except Exception as e:
            logger.warning("unknown pdf format")

    @staticmethod
    def parse_docx(file_data: bytes) -> Dict[str, Any]:
        """
        Parse DOCX file to extract text and metadata.

        Args:
            file_data: DOCX file as bytes

        Returns:
            Dictionary with 'text' and 'metadata'
        """
        try:
            docx_file = BytesIO(file_data)
            doc = Document(docx_file)

            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n".join(table_text))

            extracted_text = "\n\n".join(text_parts)

            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "format": "docx",
                "file_size": len(file_data)
            }

            # Try to get document properties
            if doc.core_properties:
                metadata["author"] = doc.core_properties.author or "Unknown"
                metadata["title"] = doc.core_properties.title or "Untitled"
                if doc.core_properties.created:
                    metadata["created"] = str(doc.core_properties.created)

            # Clean the text
            cleaned_text = DocumentParser.clean_text(extracted_text)

            return {
                "text": cleaned_text,
                "metadata": metadata
            }

        except Exception as e:
            logger.warning("unknown docx format")

    @staticmethod
    def parse_text(file_data: bytes) -> Dict[str, Any]:
        """
        Parse plain text file.

        Args:
            file_data: Text file as bytes

        Returns:
            Dictionary with 'text' and 'metadata'
        """
        try:
            # Try UTF-8 first
            text = file_data.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                text = file_data.decode('latin-1')
            except Exception as e:
                logger.warning("unknown txt format")

        # Clean the text
        cleaned_text = DocumentParser.clean_text(text)

        metadata = {
            "format": "txt",
            "file_size": len(file_data),
            "characters": len(cleaned_text)
        }

        return {
            "text": cleaned_text,
            "metadata": metadata
        }

    @staticmethod
    def parse_markdown(file_data: bytes) -> Dict[str, Any]:
        """
        Parse Markdown file.

        Args:
            file_data: Markdown file as bytes

        Returns:
            Dictionary with 'text' and 'metadata'
        """
        # Markdown is just text with special syntax
        result = DocumentParser.parse_text(file_data)
        result["metadata"]["format"] = "md"
        return result

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Normalize whitespace
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double

        # Remove common PDF artifacts
        text = re.sub(r'\x00', '', text)  # Null characters
        text = re.sub(r'[\x01-\x08\x0B-\x0C\x0E-\x1F]', '', text)  # Control chars

        # Clean up page breaks (but keep page markers)
        text = re.sub(r'\n\s*\n\[Page (\d+)\]\s*\n\s*\n', r'\n\n[Page \1]\n\n', text)

        return text.strip()

    @staticmethod
    def parse(file_data: bytes, file_type: str) -> Dict[str, Any]:
        """
        Parse document based on file type.

        Args:
            file_data: File content as bytes
            file_type: File extension (.pdf, .docx, .txt, .md)

        Returns:
            Dictionary with 'text' and 'metadata'

        Raises:
            ValueError: If file type is unsupported or parsing fails
        """
        # Map file types to parsers
        parsers = {
            ".pdf": DocumentParser.parse_pdf,
            "pdf": DocumentParser.parse_pdf,
            ".docx": DocumentParser.parse_docx,
            "docx": DocumentParser.parse_docx,
            ".txt": DocumentParser.parse_text,
            "txt": DocumentParser.parse_text,
            ".md": DocumentParser.parse_markdown,
            "md": DocumentParser.parse_markdown
        }

        # Normalize file type
        file_type = file_type.lower().lstrip()

        # Get parser
        parser = parsers.get(f"{file_type}")
        supported_types = list(set(k.lstrip('.') for k in parsers.keys()))
        if not parser:
            logger.warning("unknown format")

        # Parse
        result = parser(file_data)

        # Validate extraction
        if not result["text"] or len(result["text"].strip()) < 10:
            logger.warning("insufficient file content")

        return result
